import anthropic
import os
import re
from openpyxl import load_workbook
from dotenv import load_dotenv
from datetime import datetime
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

load_dotenv()

client = anthropic.Anthropic(api_key=os.environ.get("ANTHROPIC_API_KEY"))

# ==========================================
# 1. Excelから設定値を読み込む
# ==========================================
wb = load_workbook("food_prices.xlsx")
ws = wb["週次設定"]

設定 = {}
for row in ws.iter_rows(min_row=2, values_only=True):
    if row[0] is None:
        continue
    設定[row[0]] = row[1]

def 日付整形(値):
    if isinstance(値, datetime):
        return f"{値.year}年{値.month}月{値.day}日"
    return str(値)

def 前週比整形(値):
    if isinstance(値, (int, float)):
        if 値 > 0:
            return f"+{int(値)}"
        elif 値 < 0:
            return f"{int(値)}"
        else:
            return "±0"
    return str(値)

期間開始 = 日付整形(設定["期間開始日"])
期間終了 = 日付整形(設定["期間終了日"])
発表日 = 日付整形(設定["玉子発表日"])
東京前週比 = 前週比整形(設定["東京前週比"])
大阪前週比 = 前週比整形(設定["大阪前週比"])
名古屋前週比 = 前週比整形(設定["名古屋前週比"])

# ==========================================
# 2. プロンプトを動的に組み立てる
# ==========================================
prompt = f"""以下の情報を収集し、【出力フォーマット】の形式でまとめてください。
重要：思考過程や検索プロセスの説明は一切不要です。指定されたフォーマットの本文のみを出力してください。

【対象期間】
{期間開始}〜{期間終了}

【収集項目】

① 玉子価格（既に取得済み）
　東京 基準値：{設定["東京玉子基準値"]}円/kg　前週比：{東京前週比}
　大阪 基準値：{設定["大阪玉子基準値"]}円/kg　前週比：{大阪前週比}
　名古屋 基準値：{設定["名古屋玉子基準値"]}円/kg　前週比：{名古屋前週比}
　市況：{設定["玉子市況"]}　発表日：{発表日}

② 野菜価格（自動収集）
農林水産省の野菜卸売価格情報を検索し、対象期間最終日時点の以下13品目の前年比（%）・主産地・状況を取得してください。
品目：キャベツ、はくさい、レタス、ほうれんそう、ねぎ、きゅうり、トマト、なす、ピーマン、だいこん、にんじん、じゃがいも、たまねぎ

③ 食品値上げ関連報道（自動収集）
対象期間内に報道された食品・飲料の値上げニュースを検索してください。
※ 対象期間外の報道は含めないでください

④ 厨房消耗品の値上げ関連報道（自動収集）
対象期間内に報道された業務用消耗品（アルミホイル、クッキングシート、使い捨て手袋、洗剤、食品トレー、ラップ、ゴミ袋など）の値上げニュースを検索してください。
※ 対象期間外の報道は含めないでください

⑤ 食中毒情報（自動収集）
対象期間内に発生・報道された食中毒事例を検索してください。
※ 対象期間外の事例は含めないでください

【出力フォーマット】
冒頭に説明文を入れず、必ず「━━━」で始めてください。

━━━━━━━━━━━━━━━━━━━━━━
【食品動向】{期間開始}〜{期間終了}
━━━━━━━━━━━━━━━━━━━━━━

●玉子価格（卸値先週比）Mサイズ ※JA全農たまご
東京{設定["東京玉子基準値"]}円（前週比{東京前週比}）、大阪{設定["大阪玉子基準値"]}円（前週比{大阪前週比}）、名古屋{設定["名古屋玉子基準値"]}円（前週比{名古屋前週比}）　基準値 {発表日}発表
市況：{設定["玉子市況"]}

●野菜価格（卸値前年比）
品目 | 前年比 | 主産地 | 状況・背景
キャベツ | 000% | ◯◯県 | ◯◯◯◯◯
（13品目分）

━━━━━━━━━━━━━━━━━━━━━━
【記事(値上げ他)】
━━━━━━━━━━━━━━━━━━━━━━

●食品値上げ関連報道一覧
報道日 | 食品名・カテゴリ | 値上げ時期 | 値上げ幅 | 報道元
（対象期間内のみ。なければ「対象期間内の報道なし」）

●厨房消耗品の値上げ関連報道一覧
報道日 | 品目名 | 時期 | 値上げ幅 | 報道元
（対象期間内のみ。なければ「対象期間内の報道なし」）

●食中毒・その他
発生日 | 原因 | 発生場所・患者数 | 対応
（対象期間内のみ。なければ「対象期間内の報告なし」）

━━━━━━━━━━━━━━━━━━━━━━
【総括】
━━━━━━━━━━━━━━━━━━━━━━
（200字程度。卵・野菜価格動向、食品値上げ、食中毒情報を盛り込む）
"""

# ==========================================
# 3. Claude APIを呼び出す
# ==========================================
print(f"=== {期間開始}〜{期間終了} の週次レポートを生成中 ===")
print("Web検索を実行しています。1〜3分かかります...\n")

message = client.messages.create(
    model="claude-opus-4-5",
    max_tokens=8000,
    tools=[{
        "type": "web_search_20250305",
        "name": "web_search",
        "max_uses": 15,
    }],
    messages=[{"role": "user", "content": prompt}]
)

# ==========================================
# 4. テキスト抽出 + メタコメント除去
# ==========================================
出力テキスト = ""
for block in message.content:
    if block.type == "text":
        出力テキスト += block.text

match = re.search(r"━━━", 出力テキスト)
if match:
    出力テキスト = 出力テキスト[match.start():]

print("=== 生成された報告書 ===\n")
print(出力テキスト)

# ==========================================
# 5. テキストファイルに保存
# ==========================================
タイムスタンプ = datetime.now().strftime('%Y%m%d_%H%M%S')
txt_ファイル名 = f"週次報告書_{タイムスタンプ}.txt"
with open(txt_ファイル名, "w", encoding="utf-8") as f:
    f.write(出力テキスト)

# ==========================================
# 6. Word文書(.docx)に保存（表組み対応版）
# ==========================================

def セルを薄いグレーで塗る(セル):
    """セルの背景を薄いグレーにする（ヘッダー行用）"""
    tcPr = セル._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), 'D9E2F3')  # 薄い青グレー
    tcPr.append(shd)

doc = Document()

# タイトル
title = doc.add_heading(f"食品動向 週次報告書", level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# 期間
期間段落 = doc.add_paragraph()
期間段落.alignment = WD_ALIGN_PARAGRAPH.CENTER
期間ラン = 期間段落.add_run(f"対象期間：{期間開始}〜{期間終了}")
期間ラン.font.size = Pt(11)
期間ラン.bold = True

doc.add_paragraph()  # 空行

# 行を一行ずつ処理する状態管理
行リスト = 出力テキスト.split("\n")
i = 0

while i < len(行リスト):
    line = 行リスト[i].rstrip()
    
    # 空行
    if not line:
        doc.add_paragraph()
        i += 1
        continue
    
    # 区切り線はスキップ
    if line.startswith("━━━"):
        i += 1
        continue
    
    # 大見出し【〜】
    if line.startswith("【") and line.endswith("】"):
        doc.add_heading(line, level=1)
        i += 1
        continue
    
    # 期間付き見出し【食品動向】yyyy/m/d〜yyyy/m/d
    if "【" in line and "】" in line:
        doc.add_heading(line, level=1)
        i += 1
        continue
    
    # 中見出し●
    if line.startswith("●"):
        doc.add_heading(line, level=2)
        i += 1
        continue
    
    # 「|」が含まれる行=表のデータ
    if "|" in line:
        # ここから連続する表行を集める
        表行群 = []
        while i < len(行リスト):
            current = 行リスト[i].rstrip()
            if "|" in current:
                # 「|」で分割し、各セルの前後空白を除去
                セル群 = [cell.strip() for cell in current.split("|")]
                表行群.append(セル群)
                i += 1
            else:
                break
        
        # 表を作成
        if 表行群:
            列数 = max(len(行) for 行 in 表行群)
            # 全行を列数に揃える（不足分は空文字）
            for 行 in 表行群:
                while len(行) < 列数:
                    行.append("")
            
            表 = doc.add_table(rows=len(表行群), cols=列数)
            表.style = "Light Grid Accent 1"  # Word標準スタイル
            表.alignment = WD_TABLE_ALIGNMENT.CENTER
            
            # データ書き込み
            for r, 行データ in enumerate(表行群):
                for c, セル値 in enumerate(行データ):
                    cell = 表.cell(r, c)
                    cell.text = セル値
                    # 1行目（ヘッダー）は太字+背景色
                    if r == 0:
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                run.bold = True
                        セルを薄いグレーで塗る(cell)
            
            doc.add_paragraph()  # 表の後に空行
        continue
    
    # 通常段落
    doc.add_paragraph(line)
    i += 1

docx_ファイル名 = f"週次報告書_{タイムスタンプ}.docx"
doc.save(docx_ファイル名)

print(f"\n=== {txt_ファイル名} に保存しました ===")
print(f"=== {docx_ファイル名} に保存しました ===")